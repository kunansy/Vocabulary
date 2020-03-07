from sys import argv
import comtypes.client
from docx import Document
from docx.shared import Pt
from hashlib import sha3_512
from functools import reduce

from itertools import groupby
from datetime import datetime
from xlrd import open_workbook
from xlsxwriter import Workbook
from datetime import date as DATE
from os import system, getcwd

from random import sample as SAMPLE
from random import choice as CHOICE
from random import shuffle as SHUFFLE

from constants import *
from common_funcs import *
from backup_setup import backup
from examples import SelfExamples, CorpusExamples


from PyQt5 import uic
from PyQt5.QtWidgets import QWidget, QMainWindow, QApplication


def init_vocabulary_from_file(f_name=DATA_PATH):
    # TODO: Удалить после перехода к db
    assert file_exist(f_name), \
        f"Wrong file: '{f_name}', func – init_vocabulary_from_file"

    with open(f_name, 'r', encoding='utf-8') as file:
        content = file.readlines()
        dates = map(lambda x: str_to_date(x).strftime(DATEFORMAT),
                    filter(lambda x: x.startswith('[') and '/' not in x, content))

        begin = lambda elem: content.index(f"[{elem}]\n")
        end = lambda elem: content.index(f"[/{elem}]\n")

        words_per_day = [WordsPerDay(map(Word, content[begin(i) + 1: end(i)]), i) for i in dates]
    file.close()
    return words_per_day


def first_rus_index(item: str):
    # TODO: Удалить после перехода к db
    rus_aplphabet = "абвгдеёжзийклмнопрстуфхцчшщъыьэюя"
    return list(map(lambda x: x in rus_aplphabet, item)).index(True)


def up_word(string: str, item: str):
    """ Поднять регистр слов в строке,
        еоторые содержат item
    """
    assert isinstance(item, str) and item, \
        f"Wrong item to up: '{item}', func – up_word"
    assert isinstance(string, str) and string, \
        f"Wrong string: '{string}', func – up_word"
    word = item.lower().strip()

    if word not in string.lower():
        return string

    return change_words(string, item, str.upper)


def up_word_in_def(string: str, item: str):
    """ Поднять регистр строки до ' – ',
        где есть item или которые есть в item
    """
    assert isinstance(item, str) and item, \
        f"Wrong item to up: '{item}', func – up_word_in_def"
    assert isinstance(string, str) and len(string) >= len(item) and '–' in string, \
        f"Wrong string: '{string}', func – up_word_in_def"

    word = item.lower().strip()

    if word not in string[string.index('–'):].lower():
        return string

    return string[:string.index('–')] + up_word(string[string.index('–'):], word)


def american_spelling(word: str):
    """ Соответствует ли слово американской манере письма """
    assert isinstance(word, str) and word, \
        f"Wrong word: '{word}', func – american_spelling"
    # wrong ends
    pattern = re.compile(f'(uo|tre|nce|mm|ue|se|ll|re)\W', re.IGNORECASE)
    return bool(pattern.findall(word))


def parse_str(string: str):
    """ Разбирает строку на: термин, транксрицпию,
        свойства, английское определение, русское
    """
    # TODO: Удалить после перехода к db
    word_properties, other = string.split(' – ')

    if '[' in word_properties:
        prop_begin = word_properties.index('[')
        word, properties = word_properties[:prop_begin], Properties(word_properties[prop_begin:])
    else:
        word = word_properties
        properties = Properties('')

    if word.count("|") == 2:
        trans_begin = word.index("|")
        word, transcription = word[:trans_begin], word[trans_begin:]

    else:
        transcription = ''

    if other:
        if is_russian(other):
            fri = first_rus_index(other)
            original = other[:fri].split(';')
            native = other[fri:].split(';')
        else:
            original = other.split(';')
            native = []
    else:
        original = []
        native = []

    if any(word.lower() in i.split() for i in original):
        print(f"'{word.capitalize()}' is situated in the definition too. It is recommended to replace it on ***")

    if any('.' in i for i in original):
        print(f"There is wrong symbol in the definition of the word '{word}'")

    return word, transcription, properties, original, native


def init_from_xlsx(f_name, out='tmp',
                   date=None):
    """ Преобразовать xlsx файл в список объектов класса Word,
        вывести их в файл с дополнением старого содержимого
    """
    # TODO: Редактировать после перехода к db
    assert file_exist(f_name), \
        f"File: '{f_name}' does not exist, func – init_from_xlsx"

    if not file_exist(out):
        with open(out, 'w'): pass

    if date is None:
        date = get_current_date()
    else:
        date = str_to_date(date).strftime(DATEFORMAT)

    assert f"[{date}]\n" not in open(out, 'r', encoding='utf-8').readlines(), \
        f"Date '{date}' currently exists in the '{out}' file, func – init_from_xlsx"

    rb = open_workbook(f_name)
    sheet = rb.sheet_by_index(0)

    # удаляется заглавие, введение и прочие некорректные данные
    content = list(filter(lambda x: len(x[0]) and (len(x[2]) or len(x[3])),
                          [sheet.row_values(i) for i in range(sheet.nrows)]))
    content.sort(key=lambda x: x[0])

    # группировка одинаковых слов вместе
    content = groupby(map(lambda x: Word(word=x[0], original_def=x[2], native_def=x[3]), content),
                      key=lambda x: (x.word, x.properties))

    # суммирование одинаковых слов в один объект класса Word
    result = [reduce(lambda res, elem: res + elem, list(group[1]), Word('')) for group in content]

    with open(out, 'a', encoding='utf-8') as f:
        f.write(f"\n\n[{date}]\n")
        f.write('\n'.join(map(str, result)))
        f.write(f"\n[/{date}]\n")


def create_docx(content, f_name=None,
                header='General'):
    """ Создать docx с контентом (str или Word),
        умолчательное имя файла – заголовок, 'General';
        в качестве заголовка стоит передавать date_range();
        если файл с таким именем уже существует – не создавать новый;
    """
    assert isinstance(content, list) and content, \
        f"Wrong content: '{content}', list expected, func – create_docx"
    assert isinstance(header, str) and header, \
        f"Wrong header: '{header}', str expected, func – create_docx"

    f_name = f"{header}.docx" if f_name is None else add_ext(f_name, 'docx')

    if file_exist(f"{DOC_FOLDER}\\{f_name}"):
        print(f"Document named '{f_name}' still exist")
        return

    docx_document = Document()

    # Шрифт и его размер
    doc_style = docx_document.styles['Normal']
    font = doc_style.font
    font.name = 'Avenir Next Cyr'
    font.size = Pt(16)

    docx_document.add_heading(f"{header}", 0)

    for num, word in enumerate(content, 1):
        new_paragraph = docx_document.add_paragraph()
        new_paragraph.style = doc_style

        new_paragraph.add_run(f"{num}. ").bold = True

        if isinstance(word, str):
            new_paragraph.add_run(f"{word[0].upper()}{word[1:]}")
        elif isinstance(word, Word):
            new_paragraph.add_run(f"{word.word}".capitalize()).bold = True
            new_paragraph.add_run(f" – {word.get_original(def_only=True)} {word.get_native(def_only=True)}")

    docx_document.save(f"{DOC_FOLDER}\\{f_name}")


def create_pdf(content: list, f_name: str):
    """ Созать pdf-файл с контентом (str или Word) через docx-файл:
        создать docx-посредника преобразовать к pdf;
        если файл с таким именем уже существует – не создавать новый;
    """
    assert isinstance(content, list) and content, \
        f"Wrong content: '{content}', list expected, func – create_pdf"
    assert isinstance(f_name, str) and f_name, \
        f"Wrong f_name: '{f_name}', str expected, func – create_pdf"

    if file_exist(f"{PDF_FOLDER}\\{f_name}", 'pdf'):
        print(f"PDF-file named '{f_name}' still exist")
        return
    # файл-посредник
    mediator = f"temp_{len(content)}_{get_current_date(DATEFORMAT)}.docx"
    create_docx(content, f_name=mediator)

    f_name = add_ext(f_name, 'pdf')

    microsoft_word_client = comtypes.client.CreateObject('Word.Application')
    mediator = microsoft_word_client.Documents.Open(f"{getcwd()}\\{DOC_FOLDER}\\{mediator}")
    mediator.SaveAs(f"{getcwd()}\\{PDF_FOLDER}\\{f_name}", FileFormat=17)

    mediator.Close()
    microsoft_word_client.Quit()

    system(f'del "{getcwd()}\\{DOC_FOLDER}\\{f_name}"')


def word_id(item):
    """ Получить ID переданного строкой или Word-объектом слова
        ID – первые и последние восемь символов sha3_512 хеша этого слова
    """
    assert isinstance(item, str) or isinstance(item, Word), f"Wrong word: '{item}', func – word_id"
    # пустой item – пустой ID
    if not item:
        return ''

    word = item if isinstance(item, str) else item.word

    _id = sha3_512(bytes(word, encoding='utf-8')).hexdigest()
    return _id[:ID_LENGTH//2] + _id[-ID_LENGTH//2:]


def search_by_attribute(sample, item):
    """ Ищет в выборке объект класса Word, один из
        атрибутов которого соответствует искомому айтему
    """
    assert isinstance(sample, list) and sample and all(isinstance(i, Word) for i in sample), \
        f"Wrong words_list: '{sample}', func – search_by_attribute"
    assert isinstance(item, str) or isinstance(item, Word), \
        f"Wrong item: '{item}', func – search_by_attribute"

    item = item.lower().strip() if isinstance(item, str) else item.word

    # Если в переданном айтеме есть русский символ – соответствие
    # может быть только с русским определением
    if is_russian(item):
        try:
            return list(filter(lambda x: x.get_native(def_only=True).lower() == item, sample))[0]
        except:
            print(f"Not found word, which attributes fit with item: '{item}' in the sample: '{sample}'")
            return []

    # Поиск соответствия в самих словах
    in_word = list(filter(lambda x: x.word == item, sample))
    try:
        return in_word[0] if len(in_word) else list(filter(lambda x: x.get_original(def_only=True).lower() == item, sample))[0]
    except:
        return f"Word, which attributes fit with: '{item}' in the sample: '{sample}', \nnot found"


def get_current_date(dateformat=DATEFORMAT):
    """ Current date with str/DATE: format is None –
        DATE, else – dateformat (default – dd.mm.yy)
    """
    return datetime.now().date() if dateformat is None else datetime.now().strftime(dateformat)


def diff_words(base):
    """ Most difficult words, from json base """
    ids = diff_words_id()
    return base.search_by_id(ids)


def backup_repeat_log():
    """ Backup лога повторений """
    print("Repeat log backupping...")
    backup(file_name(REPEAT_LOG_PATH), REPEAT_LOG_PATH)


class RepeatWords(QMainWindow):
    def __init__(self, words, mode=1,
                 window_title='Repeat'):
        assert file_exist(MAIN_WINDOW_PATH), \
            "Main window does not exist, func – RepeatWords.__init__"
        assert file_exist(EXAMPLES_WINDOW_PATH), \
            "Examples window does not exist, func – RepeatWords.__init__"
        assert file_exist(MESSAGE_WINDOW_PATH), \
            "Message window does not exist, func – RepeatWords.__init__"
        assert file_exist(SHOW_WINDOW_PATH), \
            "Show window does not exist, func – RepeatWords.__init__"

        assert isinstance(words, list) and words and all(isinstance(i, Word) for i in words), \
            f"Wrong words: '{words}', func – RepeatWords.__init__"
        assert (isinstance(mode, int) and mode in range(len(REPEAT_MODS))) or \
               (isinstance(mode, str) and mode in RepeatWords), \
            f"Wrong mode: '{mode}', func – RepeatWords.__init__"
        assert isinstance(window_title, str), \
            f"Wrong window title: '{window_title}', func – RepeatWords.__init__"

        super().__init__()
        uic.loadUi(MAIN_WINDOW_PATH, self)
        self.initUI(window_title)

        self.word = Word('')
        # self.repeating_word_index = None
        # self.repeated_words_indexes = []
        SHUFFLE(words)
        self.words = words[:]
        self.wrong_translations = list(reversed(words[:]))

        self.mode = None
        self.are_you_right = None
        self.init_button = None
        self.main_item = None

        if isinstance(mode, int):
            self.mode = mode
        elif isinstance(mode, str):
            self.mode = REPEAT_MODS[mode]

        self.init_fit_mode()

        if not file_exist(REPEAT_LOG_PATH):
            with open(REPEAT_LOG_PATH, 'w', encoding='utf-8'): pass

    def initUI(self, window_title):
        self.ExamplesWindow = ExamplesWindow(self, [])
        self.MessageWindow = MessageWindow(self, [])
        self.ShowWindow = ShowWindow(self, [])
        self.setWindowTitle(window_title)

        # кнопки выбора вариантов
        self.choice_buttons = [
            self.ChoiceButton1,
            self.ChoiceButton2,
            self.ChoiceButton3,
            self.ChoiceButton4,
            self.ChoiceButton5,
            self.ChoiceButton6
        ]
        [self.choice_buttons[i].clicked.connect(self.are_you_right)
         for i in range(len(self.choice_buttons))]

        self.ExitButton.clicked.connect(self.close)
        self.ExitButton.setText('Exit')

        self.HintButton.clicked.connect(self.hint)
        self.HintButton.setText('Hint')

        self.ShowButton.clicked.connect(self.show_words)
        self.ShowButton.setText('Show')

    def init_fit_mode(self):
        """ Работа в соответствии с mode """
        self.are_you_right = lambda x: x.lower() == self.init_button(self.word).lower()

        if self.mode == 1:
            self.init_button = lambda x: x.get_native(def_only=True).capitalize()
            self.main_item = lambda: self.word.word.capitalize()
        elif self.mode == 2:
            self.init_button = lambda x: x.word.capitalize()
            self.main_item = lambda: self.word.get_native(def_only=True).capitalize()
        elif self.mode == 3:
            self.init_button = lambda x: x.get_native(def_only=True).capitalize()
            self.main_item = lambda: self.word.get_original(def_only=True).capitalize()
        elif self.mode == 4:
            self.init_button = lambda x: x.get_original(def_only=True).capitalize()
            self.main_item = lambda: self.word.get_native(def_only=True).capitalize()

    def test(self):
        # чтобы не остаться с пустым списком ошибочных переводов под конец
        if self.word and self.word not in self.wrong_translations:
            self.wrong_translations.append(self.word)

        if len(self.words) == 1:
            self.WordsRemainLabel.setText("The last one")
        else:
            self.WordsRemainLabel.setText(f"Remain {len(self.words)} words")

        self.word = CHOICE(self.words)
        # self.repeating_word_index = CHOICE(range(len(self.words)))
        # self.repeating_word = CHOICE([iter(i) for i in self.words])
        self.words.remove(self.word)
        self.wrong_translations.remove(self.word)

        self.WordToReapeatBrowser.setText(self.main_item())

        self.set_buttons()

    def set_buttons(self):
        [self.choice_buttons[i].setText('')
         for i in range(len(self.choice_buttons))]

        # ставится ли верное определение
        is_right_def = True

        wrong_translations = SAMPLE(self.wrong_translations, len(self.choice_buttons))

        for i in SAMPLE(range(len(self.choice_buttons)), len(self.choice_buttons)):
            w_item = CHOICE(wrong_translations)
            wrong_translations.remove(w_item)

            if is_right_def:
                self.choice_buttons[i].setText(self.init_button(self.word))
                is_right_def = False
            else:
                self.choice_buttons[i].setText(self.init_button(w_item))

    # def next(self):
    #     if self.repeating_word_index < len(self.words):
    #         self.repeating_word_index += 1
    #         self.text()
    #
    # def past(self):
    #     if self.repeating_word_index > 0:
    #         self.repeating_word_index -= 1
    #         self.test()

    def show_result(self, result, style):
        self.MessageWindow.display(
            message=result, style=style)

    def are_you_right(self):
        if self.are_you_right(self.sender().text()):
            self.show_result(result='<i>Excellent</i>',
                             style="color: 'green';")
            if len(self.words) > 0:
                self.test()
            else:
                self.close()
        else:
            self.log(w_choice=self.sender().text(),
                     item=self.word)

            self.show_result(result='<b>Wrong</b>',
                             style="color: 'red';")

    def hint(self):
        """ Показать подсказку: связные слова и примеры """
        if self.mode != 1:
            return

        self.ExamplesWindow.display(word=self.word.word,
                                    message=f"{self.HintButton.text()}",
                                    style="color: blue")

    def show_words(self):
        self.ShowWindow.display(items=self.words,
                                window_title=self.windowTitle())

    def show(self):
        super().show()
        self.show_words()

    def close(self):
        """ По нажатии на кнопку 'Exit' закрыть все окна """
        self.ExamplesWindow.close()
        self.MessageWindow.close()
        self.ShowWindow.close()
        super().close()

    def log(self, item, w_choice):
        """ Логгирование в файл ошибочных вариантов в json:
            {ID_слова: {ID ошибочного варианта: количество ошибок}}
        """
        assert file_exist(REPEAT_LOG_PATH), \
            "Wrong log file, func – RepeatWords.log"
        assert isinstance(item, str) or isinstance(item, Word), \
            f"Wrong word: '{item}', func – RepeatWords.log"
        assert isinstance(w_choice, str) and w_choice, \
            f"Wrong w_choice: '{w_choice}', func – RepeatWords.log"

        repeating_word_id = item.get_id() if isinstance(item, Word) else word_id(item)

        # Найти ID слова по выбранному варианту:
        wrong_choice_id = search_by_attribute(self.wrong_translations, w_choice).get_id()

        data = load_json_dict(REPEAT_LOG_PATH)

        if repeating_word_id in data:
            if wrong_choice_id in data[repeating_word_id]:
                data[repeating_word_id][wrong_choice_id] += 1
            else:
                data[repeating_word_id][wrong_choice_id] = 1
        else:
            data[repeating_word_id] = {wrong_choice_id: 1}

        dump_json_dict(data=data, filename=REPEAT_LOG_PATH)


class ExamplesWindow(QWidget):
    def __init__(self, *args):
        super().__init__()
        uic.loadUi(EXAMPLES_WINDOW_PATH, self)

        self.s_examples_base = SelfExamples()
        self.c_examples_base = CorpusExamples()

        self.word = ''
        self.j_word = ''

        self.c_examples = []
        self.s_examples = []
        self.linked_words = []

        # жирный курсив
        self.bi = lambda x: f"<b><i>{x}</i></b>".upper()
        self.initUI()

    def initUI(self):
        self.setWindowTitle('Examples')

        self.ResultLabel.setText('')
        self.ExamplesBrowser.setText('')

        self.CorpusExamplesRButton.setChecked(True)

        self.CorpusExamplesRButton.clicked.connect(self.set_corpus_examples)
        self.SelfExamplesRButton.clicked.connect(self.set_self_examples)
        self.LinkedWordsRButton.clicked.connect(self.set_linked_words)

    def get_linked_words(self):
        """ Инициализировать linked_words,
            получив связные слова
        """
        try:
            # для запроса нужно само слово
            # без предлогов и прочего
            content = get_synonyms(self.j_word)
        except:
            content = []
        self.linked_words = list(map(str.capitalize, content))

    def get_corpus_examples(self):
        """ Инициализировать self.c_examples примерами к
            слову из глобальной корпусной базы
        """
        # если примеров мало – запросить ещё
        if self.c_examples_base.count(self.j_word) < 6:
            self.c_examples_base.new_examples(self.j_word)

        if self.c_examples_base.count(self.j_word) > 0:
            # показать только оригиналы,
            # отсортированные по длине
            self.c_examples = list(map(lambda x: x[0], self.c_examples_base(self.j_word)))
            self.c_examples.sort(key=len)

    def get_self_examples(self):
        """ Инициализировать self.s_examples
            примерами к слову
        """
        self.s_examples = self.s_examples_base.find(self.j_word)

    def examples_exist(self):
        """ Существуют ли хоть какие-то примеры """
        return self.c_examples or \
               self.s_examples or \
               self.linked_words

    def set_corpus_examples(self):
        """ Установить в поле 'примеры' корпусные примеры,
            либо 'Corpus examples not found'
        """
        if self.c_examples:
            self.show_examples(self.c_examples)
        else:
            self.show_examples([f"Corpus examples not '{self.word}' found"])

    def set_self_examples(self):
        """ Установить в поле 'примеры' свои примеры,
            либо 'Self examples not found'
        """
        if self.s_examples:
            self.show_examples(self.s_examples)
        else:
            self.show_examples([f"Self examples to '{self.word}' not found"])

    def set_linked_words(self):
        """ Установить в поле 'примеры' связные слова """
        if self.linked_words:
            self.show_examples(self.linked_words)
        self.show_examples([f"Linked words to '{self.word}' not found"])

    def show_examples(self, examples):
        examples = map(lambda x: f"{self.bi(x[0])}. {change_words(x[1], self.j_word, self.bi)}<br>",
                       enumerate(examples, 1))

        self.ExamplesBrowser.setText('\n'.join(examples))

    def set_examples(self):
        if self.CorpusExamplesRButton.isChecked():
            self.set_corpus_examples()
        elif self.SelfExamplesRButton.isChecked():
            self.set_self_examples()
        elif self.LinkedWordsRButton.isChecked():
            self.set_linked_words()

    def display(self, word, message, style=''):
        assert isinstance(word, str) and word, \
            f"Wrong word: '{word}', str expected, func – Examples.display"
        assert isinstance(message, str) and message, \
            f"Wrong result: '{message}', func – Examples.display"

        if word != self.word:
            self.c_examples = []
            self.s_examples = []
            self.linked_words = []

            self.word = word
            self.j_word = just_word(word)

            self.get_linked_words()
            self.get_corpus_examples()
            self.get_self_examples()

        # вывод корпусных примеров по умолчанию
        self.CorpusExamplesRButton.setChecked(True)

        if not self.examples_exist():
            # если примероы нет – не открывать окно
            self.close()
            return

        if style:
            self.ResultLabel.setStyleSheet(style)
        self.ResultLabel.setText(message)
        self.set_examples()
        self.show()


class MessageWindow(QWidget):
    def __init__(self, *args):
        """ Окно отображения сообщений, правильности или
            неправильности ответов
        """
        super().__init__()
        uic.loadUi(MESSAGE_WINDOW_PATH, self)

        self.initUI()

    def initUI(self):
        self.setWindowTitle('Message')
        self.MessageText.setText('')

    def display(self, message, style=''):
        assert isinstance(message, str) and message, \
            f"Wrong message: '{message}', str expected, func – Message.display"

        if style:
            self.MessageText.setStyleSheet(style)

        self.MessageText.setText(message)
        self.show()


class ShowWindow(QWidget):
    def __init__(self, *args):
        super().__init__()
        uic.loadUi(SHOW_WINDOW_PATH, self)

    def display(self, items, window_title='Show'):
        assert isinstance(items, list) and items, \
            f"Wrong items: '{items}', func – Show.display"
        assert isinstance(window_title, str) and window_title, \
            f"Wrong window_title: '{window_title}', func – Show.display"

        self.setWindowTitle(window_title)

        self.LearnedWordsBrowser.setText('\n'.join(map(
            lambda x: f"<i><b>{x.word.capitalize()}</b></i> – {x.get_native(def_only=True)}<br>",
            sorted(items))))
        self.show()


class Properties:
    def __init__(self, properties):
        assert isinstance(properties, str) or isinstance(properties, list), \
            f"Wrong properties: '{properties}', func – Properties.__init__"

        if isinstance(properties, str):
            properties = properties.replace('[', '').replace(']', '').split(',')
        properties = list(filter(len, map(lambda x: x.strip().lower(), properties)))

        self.properties = properties[:]

    def __eq__(self, other):
        if isinstance(other, Properties):
            return sorted(self.properties) == sorted(other.properties)
        return other.lower() in self.properties

    def __ne__(self, other):
        return not (self == other)

    def __getitem__(self, item):
        return self == item

    def __getattr__(self, item):
        return self == item

    def __len__(self):
        return len(self.properties)

    def __add__(self, other):
        if isinstance(other, Properties):
            return Properties(self.properties + other.properties)

    def __hash__(self):
        return hash(tuple(self.properties))

    def __str__(self):
        return f"[{', '.join(map(str.capitalize, self.properties))}]"


class Word:
    def __init__(self, word='', transcription='',
                 properties='', original_def=[], native_def=[]):
        """
        :param word: original word to learn
        :param properties: POS, language level, formal, ancient,
        if verb (transitivity, ) if noun (countable, )
        :param original_def: original definitions of the word: list
        :param native_def: native definitions of the word (maybe don't exist): list
        """
        assert isinstance(word, str), \
            f"Wrong word: '{word}', '{word}', func – Word.__init__"
        assert isinstance(transcription, str), \
            f"Wrong transcription: '{transcription}', func – Word.__init__"
        assert isinstance(properties, str) or isinstance(properties, Properties), \
            f"Wrong properties: '{properties}', func – Word.__init__"
        assert isinstance(original_def, list) or isinstance(original_def, str), \
            f"Wrong original_def: '{original_def}', func – Word.__init__"
        assert isinstance(native_def, list) or isinstance(native_def, str), \
            f"Wrong native_def: '{native_def}', func – Word.__init__"

        if ' – ' in word:
            self.__init__(*parse_str(word))
        else:
            self.word = word.lower().strip()
            self.id = word_id(self)
            self.transcription = transcription.replace('|', '').strip()

            if isinstance(native_def, str):
                native_def = native_def.split(';')

            if isinstance(original_def, str):
                original_def = original_def.split(';')

            self.original = list(filter(len, map(str.strip, original_def)))
            self.native = list(filter(len, map(str.strip, native_def)))

            self.properties = ''

            if isinstance(properties, Properties):
                self.properties = properties
            elif isinstance(properties, str):
                self.properties = Properties(properties)

    def get_native(self, def_only=False):
        """
        Вовзращает русские определения
        :param def_only: True – только определения, False – термин и определения
        """
        if def_only:
            return '; '.join(self.native)
        return f"{self.word} – {'; '.join(self.native)}".capitalize()

    def get_original(self, def_only=False):
        """
        Возвращает английские определения
        :param def_only: True – только определения, False – термин и определения
        """
        if def_only:
            return '; '.join(self.original)
        return f"{self.word} – {'; '.join(self.original)}".capitalize()

    def get_transcription(self):
        return f"/{self.transcription}/" * (len(self.transcription) != 0)

    def get_id(self):
        """ Вернуть либо ID слова """
        return self.id if self.id else word_id(self)

    def is_fit(self, *properties):
        """ Соотвествует ли слово переданным свойствам """
        return all(self.properties[i] for i in properties)

    def __getitem__(self, index):
        return self.word[index]

    def __iter__(self):
        return iter(self.word)

    def __add__(self, other):
        """
        :param other: строку с ' – ' для преобразования к Word или сам объект класса Word
        :return: если термины одни и те же – просуммирует списки определений и свойств
        """
        if isinstance(other, str) and ' – ' not in other or not isinstance(other, Word):
            raise ValueError(f"'Operation +' between 'class Word' and '{type(other)}' does not support")

        if isinstance(other, str):
            other = Word(other)

        if self != other and len(self) != 0 and len(other) != 0:
            raise ValueError(f"'Operator +' demands for the equality words")

        return Word(
            max(self.word, other.word),
            self.transcription,
            other.properties + self.properties,
            other.original[:] + self.original[:],
            other.native[:] + self.native[:],
        )

    def __eq__(self, other):
        if isinstance(other, str):
            return self.word == other.lower().strip()
        if isinstance(other, Word):
            return self.word == other and \
                   self.properties == other.properties
        if isinstance(other, int):
            return len(self.word) == other

    def __ne__(self, other):
        return not (self == other)

    def __gt__(self, other):
        if isinstance(other, str):
            return self.word > other.lower().strip()
        if isinstance(other, Word):
            return self.word > other.word
        if isinstance(other, int):
            return len(self.word) > other

    def __lt__(self, other):
        return self != other and not (self > other)

    def __ge__(self, other):
        return self > other or self == other

    def __le__(self, other):
        return self < other or self == other

    def __len__(self):
        return len(self.word)

    def __bool__(self):
        return bool(self.word)

    def __contains__(self, item):
        """ Если в item есть '–', или item содержит
            кириллические символы – посик по определениям
        """
        if isinstance(item, str):
            item = item.lower().strip()

            # работает по определениям
            if '–' in item or is_russian(item):
                item = item.replace('–', '').strip()

                if is_russian(item):
                    return item in self.get_native(def_only=True)
                return item in self.get_original(def_only=True)

            # по словам
            return item in self.word or self.word in item

        if isinstance(item, Word):
            return self.word in item.word or item.word in self.word

    def __str__(self):
        transcription = f" /{self.transcription}/" * (len(self.transcription) != 0)
        properties = f" {self.properties}" * (len(self.properties) != 0)
        learn = f"{'; '.join(self.original)}\t" * (len(self.original) != 0)
        rus = f"{'; '.join(self.native)}" * (len(self.native) != 0)

        return f"{self.word.capitalize()}{transcription}{properties} – {learn}{rus}"

    def __hash__(self):
        return hash(
            hash(self.word) +
            hash(self.transcription) +
            hash(self.properties) +
            hash(tuple(self.native)) +
            hash(tuple(self.original))
        )


class WordsPerDay:
    def __init__(self, content, date):
        """
        :param content: list or iterator объектов класса Word
        :param date: дата изучения
        """
        assert isinstance(date, str) or isinstance(date, DATE), \
            f"Wrong date: {date}, func – WordsPerDay.__init__"

        self.date = str_to_date(date)
        self.content = list(sorted(content))

    def native_only(self, def_only=False):
        """
        :param def_only: return words with its definitions or not
        :return: the list of the words with its native definitions, the day contains
        """
        return reduce(
            lambda res, elem: res + [elem.get_native(def_only)],
            self.content, 
            []
        )

    def original_only(self, def_only=False):
        """
        :param def_only: return words with its definitions or not
        :return: the list of the words with its original definitions, the day contains
        """
        return reduce(
            lambda res, elem: res + [elem.get_original(def_only)],
            self.content, 
            []
        )

    def get_words_list(self, with_original=False,
                       with_rus=False):
        """
        :param with_original: return words with original its definitions or not
        :param with_rus: return words with its native definitions or not
        :return: the list of the words, the day contains, with its original/native definitions
        """
        rus = lambda x: f"{x.get_native(def_only=True)}" * int(with_rus)
        original = lambda x: f"{x.get_original(def_only=True)}" * int(with_original)

        devis = lambda x: ' – ' * ((len(rus(x)) + len(original(x))) != 0)

        value = lambda x: f"{x.word}{devis(x)}{original(x)}{rus(x)}"

        return reduce(
            lambda res, elem: res + [value(elem)], 
            self.content,
            []
        )

    def get_content(self):
        return self.content[:]

    def get_date(self, dateformat=None):
        return self.date if dateformat is None else self.date.strftime(dateformat)

    def get_info(self):
        return f"{self.get_date(DATEFORMAT)}\n{len(self)}"

    def repeat(self, **params):
        pass
        # app = QApplication(argv)
        #
        # repeat = RepeatWords(words=self.content,
        #                      window_title=self.get_date(DATEFORMAT),
        #                      **params)
        # repeat.test()
        # repeat.show()
        #
        # exit(app.exec_())

    def create_docx(self):
        create_docx(self.content,
                    header=self.get_date(DATEFORMAT))

    def create_pdf(self):
        create_pdf(self.content,
                   f_name=self.get_date(DATEFORMAT))

    def search_by_properties(self, *properties):
        return list(filter(lambda x: x.is_fit(*properties), self.content))

    def __len__(self):
        """ Количество слов """
        return len(self.content)

    def __str__(self):
        date = f"{self.get_date(DATEFORMAT)}\n"
        if len(self.content) == 0:
            return date + 'Empty'
        return date + '\n'.join(map(str, self.content))

    def __contains__(self, item):
        """ word or its def """
        assert isinstance(item, str) or isinstance(item, Word), \
            f"Wrong item: '{type(item)}', func – WordsPerDay.__contains__"

        return any(item in i for i in self.content)

    def __getitem__(self, item):
        """
        :param item: word or index or slice
        :return: object Word or WordsPerDay item in case of slice
        """
        if (isinstance(item, str) or isinstance(item, Word)) and item in self:
            return list(filter(lambda x: item in x, self.content))
        if isinstance(item, int) and abs(item) <= len(self.content):
            return self.content[item]
        if isinstance(item, slice):
            return self.__class__(self.content[item], self.date)

        raise IndexError(f"Wrong index or item {item} does not exist in {self.get_date(DATEFORMAT)}")

    def __iter__(self):
        return iter(self.content)

    def __lt__(self, other):
        if isinstance(other, WordsPerDay):
            return len(self.content) < len(other.content)
        if isinstance(other, int):
            return len(self.content) < other

    def __eq__(self, other):
        if isinstance(other, WordsPerDay):
            return len(self.content) == len(other.content)
        if isinstance(other, int):
            return len(self.content) == other

    def __gt__(self, other):
        return not self < other and not self == other

    def __hash__(self):
        return hash(
            hash(tuple(self.content)) +
            hash(self.date))


class Vocabulary:
    def __init__(self, f_name=DATA_PATH):
        assert file_exist(f_name), \
            f"Wrong file: {f_name}, func – Vocabulary.__init__"
        self.list_of_days = init_vocabulary_from_file(f_name)[:]

        # имя файла с графиком динамики изучения
        self.graphic_name = f"{TABLE_FOLDER}\\Information_{self.get_date_range()}.xlsx"

    def get_pairs_date_count(self):
        """ Вернуть пары из непустых дней:
            дата – количество изученных слов
        """
        return {i.get_date(): len(i) for i in self.list_of_days}

    def max_day_info(self):
        """ Информация о дне с max количеством
            слов: дата и само количество
        """
        max_day = max(self.list_of_days)
        return f"Maximum day {max_day.get_date(DATEFORMAT)}: {len(max_day)}"

    def min_day_info(self):
        """ Информация о дне с min количеством
            слов: дата и само количество
        """
        min_day = min(self.list_of_days)
        return f"Minimum day {min_day.get_date(DATEFORMAT)}: {len(min_day)}"

    def avg_count_of_words(self):
        """ Среднее количество изученных за день слов """
        return sum(len(i) for i in self.list_of_days) // self.duration()

    def empty_days_count(self):
        """ Количество пустых дней """
        return self.duration() - len(self.list_of_days)

    def statistics(self):
        """ Статистика о словаре:
            продолжительность; среднее количество слов; всего слов изучено;
            могло бы быть изучно, но эти дни пустые (считается умножением
            количества пустых дней на среднее количество изученных в день слов);
            min/max количества изученных слов
        """
        avg_value = self.avg_count_of_words()
        empty_count = self.empty_days_count()

        avg_words_count = f"Average value = {avg_value}"
        duration = f"Duration = {self.duration()}"
        total_amount = f"Total = {len(self)}"
        would_be_total = f"Would be total = {len(self) + avg_value * empty_count}\n" \
                         f"Lost = {self.avg_count_of_words() * empty_count} items per " \
                         f"{empty_count} empty days"
        min_max = f"{self.max_day_info()}\n{self.min_day_info()}"

        return f"{duration}\n{avg_words_count}\n{total_amount}\n{would_be_total}\n\n{min_max}"

    def get_date_list(self):
        """ Список дат DATE-объектами """
        return list(map(lambda x: x.get_date(), self.list_of_days))

    def get_date_range(self):
        """ Дата первого дня-дата последнего дня """
        return f"{self.begin().strftime(DATEFORMAT)}-" \
               f"{self.end().strftime(DATEFORMAT)}"

    def get_item_before_now(self, days_count: int):
        """ Вернуть непустой день, чей индекс = len - days_count """
        assert isinstance(days_count, int) and days_count >= 0, \
            f"Wrong days_count: '{days_count}', func – Vocabulary.get_item_before_now"

        index = len(self.list_of_days) - days_count - 1

        assert index >= 0, \
            f"Wrong index: '{index}', func – Vocabulary.get_item_before_now"

        return self.list_of_days[index]

    def common_words_list(self):
        """ Вернуть отсортированный список всех слов """
        return list(sorted(reduce(
            lambda result, element: result + element.get_content(),
            self.list_of_days,
            [])))

    def duration(self):
        """ Продолжительность ведения словаря """
        return (self.end() - self.begin()).days + 1

    def create_visual_info(self):
        """ Создать Excel файл с графиком динамики изучения слов:
            по оси x – непустые дни, по y – количество изученых слов
        """
        if file_exist(self.graphic_name):
            return
        sheet_name = 'Graphic'

        workbook = Workbook(self.graphic_name)
        workbook.set_properties({
            'title': "Vocabulary learning",
            'author': "Kolobov Kirill",
            'comments': "Created with Python and XlsxWriter"
        })
        font = "Avenir Next Cyr"
        cell_format = workbook.add_format({
            'size': 16,
            'font': font,
            'align': "vcenter"
        })

        worksheet = workbook.add_worksheet(sheet_name)

        date_count = self.get_pairs_date_count()

        worksheet.set_column('A:A', 17)

        # TODO: можно ли как-то это ускорить, избежав for?
        for row, (date, count) in enumerate([(date, count) for date, count in date_count.items()]):
            worksheet.write(row, 0, date.strftime(DATEFORMAT), cell_format)
            worksheet.write(row, 1, count, cell_format)

        row += 1

        worksheet.write(row, 0, "Total:", cell_format)
        worksheet.write(row, 1, f"=SUM(B1:B{row})", cell_format)

        chart = workbook.add_chart({
            'type': "line"
        })
        # line, area, bar, column

        chart.set_title({
            'name': "Amount of the original words",
            'name_font': {
                'name': font,
                'color': "black",
                'size': 16
            },
        })

        chart.add_series({
            'values': f"={sheet_name}!B1:B{row}",
            'categories': f"={sheet_name}!A1:A{row}",
            'line': {
                'color': "orange"
            },
        })

        chart.set_legend({
            'none': True
        })

        chart.set_x_axis({
            'name': 'Days',
            'name_font': {
                'name': font,
                'italic': True,
                'size': 16
            },
            'num_font': {
                    'name': font,
                    'italic': True,
                    'size': 14
                }
        })
        chart.set_y_axis({
            'name': 'Count',
            'name_font': {
                'name': font,
                'italic': True,
                'size': 16
            },
            'num_font': {
                'name': font,
                'italic': True,
                'bold': True,
                'size': 14
            }
        })

        chart.set_size({
            'width': 1280,
            'height': 520
        })

        worksheet.insert_chart('C1', chart)

        workbook.close()

    def create_docx(self):
        """ Создать docx-файл со всемии словами словаря,
            отсортированными по алфавиту; Имя файла –
            date_range текущего словаря;
        """
        create_docx(content=self.common_words_list(),
                    f_name=self.get_date_range(),
                    header=self.get_date_range())

    def create_pdf(self):
        """ Создать pdf-файл со всемии словами словаря,
            отсортированными по алфавиту; Имя файла –
            date_range текущего словаря;
        """
        create_pdf(content=self.common_words_list(),
                   f_name=self.get_date_range())

    def search(self, item):
        """ Вернуть словарь из даты изучения в ключе и
            найденными словами из этого дня в значениях
        """
        # TODO: написать комментарий к принципу поиска, иправить его
        assert isinstance(item, str) or isinstance(item, Word), \
            f"Wrong item: '{item}', func – Vocabulary.search"
        assert len(item) > 1, \
            f"Wrong item: '{item}', func – Vocabulary.search"
        assert item in self, \
            f"Word is not in the Vocabulary '{item}', func – Vocabulary.search"

        item = item.lower().strip() if isinstance(item, str) else item

        return {i.get_date(DATEFORMAT): i[item] for i in filter(lambda x: item in x, self.list_of_days)}

    def show_graph(self):
        """ Показать график, создав его в случае отсутствия """
        if not file_exist(self.graphic_name):
            self.create_visual_info()

        system(self.graphic_name)

    def info(self):
        """ Создать xlsx файл с графиком изучения слов, вернуть информацию о словаре:
            пары: день – количество изученных слов; статистика """
        self.create_visual_info()

        return '\n'.join(map(lambda day_count: f"{day_count[0]}: {day_count[1]}", self.get_pairs_date_count().items())) + \
               f"\n{DIVIDER}\n{self.statistics()}"

    def repeat(self, *items_to_repeat, **params):
        """ Запуск повторения уникальных слов при указанном mode;
            Выбор слов для повторения:
                1. Если это int: день, чей индекс равен текущему - int_value;
                2. Если это str или DATE: день с такой датой;
                3. random=n: один либо n случайных дней;
                4. most_difficult=n: n либо все слова из лога повторений
        """
        if 'random' in items_to_repeat or 'random' in params:
            if 'random' in items_to_repeat:
                repeating_days = [CHOICE(self.list_of_days)]
            else:
                assert isinstance(params['random'], int) and params['random'] <= len(self.list_of_days), \
                    f"Wrong random value: '{params['random']}', func – Vocabulary.repeat"

                repeating_days = SAMPLE(self.list_of_days, params['random'])
                params.pop('random')
        elif 'most_difficult' in items_to_repeat or 'most_difficult' in params:
            if 'most_difficult' in items_to_repeat:
                repeating_days = [WordsPerDay(self.search_by_id(*(diff_words_id())), get_current_date())]
            else:
                count = params['most_difficult']
                assert isinstance(count, int), f"Wrong most_difficult: '{count}'"

                repeating_days = [WordsPerDay(self.search_by_id(*diff_words_id()[:count]), get_current_date())]
                params.pop('most_difficult')
        else:
            days_before_now = filter(lambda x: isinstance(x, int), items_to_repeat)
            dates = filter(lambda x: isinstance(x, str) or isinstance(x, DATE), items_to_repeat)

            days_before_now = list(map(self.get_item_before_now, days_before_now))
            dates = list(map(self.__getitem__, dates))

            repeating_days = dates + days_before_now

        assert len(repeating_days) != 0, \
            f"Wrong item to repeat, func – Vocabulary.repeat"

        repeating_days.sort(key=lambda x: x.get_date())

        if len(repeating_days) > 1 and repeating_days[0].get_date() != repeating_days[-1].get_date():
            window_title = f"{repeating_days[0].get_date()}-{repeating_days[-1].get_date()}, {len(repeating_days)} days"
        else:
            window_title = repeating_days[0].get_date(DATEFORMAT)

        # Переданные айтемы могут содержать одинаковые слова
        repeating_days = sum(list(map(WordsPerDay.get_content, repeating_days)), [])
        repeating_days = list(set(tuple(repeating_days)))

        app = QApplication(argv)

        repeat = RepeatWords(words=repeating_days,
                             window_title=window_title,
                             **params)

        repeat.test()
        repeat.show()

        exit(app.exec_())

    def begin(self):
        """ Вернуть дату первого дня DATE-объектом """
        return self.list_of_days[0].get_date()

    def end(self):
        """ Вернуть дату последнего дня DATE-объектом """
        return self.list_of_days[-1].get_date()

    def search_by_properties(self, *properties):
        """ Найти слова, удовлетворяющие переданным свойствам """
        # TODO: будет ли работать при sth.search_by_properties([prop1, prop2...])?
        return list(filter(lambda x: x.is_fit(*properties), self.common_words_list()))

    def search_by_id(self, *ids):
        """ Вернуть список слов, чьи ID равны ID переданным """
        if isinstance(ids[0], list):
            ids = sum(ids, [])

        assert all(isinstance(i, str) and len(i) == ID_LENGTH for i in ids), \
            f"Wrong ids: '{ids}', func – Vocabulary.search_by_id"

        # TODO: сортировать в соответствии с порядком id
        # id_to_word = list(map(lambda x: (x.get_id(), x), self.common_words_list()))
        # print(len(id_to_word))
        # print(list(filter(lambda x: )))

        # return list(filter(lambda x: x[0] in ids, id_to_word))

    def how_to_say_in_native(self):
        """ Вернуть только слова на изучаемом языке """
        return list(map(lambda x: x.word, self.common_words_list()))

    def how_to_say_in_original(self):
        """ Вернуть только нативные определения слов """
        return reduce(lambda res, elem: res + elem.native_only(def_only=True),
                      self.list_of_days, [])

    def backup(self):
        """ Backup главного словаря """
        # добавить в имя файла промежуток дат
        f_name = add_to_file_name(file_name(DATA_PATH),
                                  f"_{self.get_date_range()}")
        print("Main data backupping...")
        backup(f_name, DATA_PATH)

    def __contains__(self, item):
        """ Есть ли слово (str или Word) или WordsPerDay
            с такой датой (только DATE) в словаре
        """
        if isinstance(item, str):
            return any(item.lower().strip() in i for i in self.list_of_days)
        if isinstance(item, Word):
            return any(item in i for i in self.list_of_days)
        if isinstance(item, DATE):
            if not any(i.date == item for i in self.list_of_days):
                # если нет равенства даты – проверка на вхождение даты в промежуток
                # [начало словаря; конец], и если дата попала в промежуток, то это
                # означает, что день с такой датой есть в словаре, но он пуст
                if self.begin() <= item <= self.end():
                    return True
                return False
            return True

    def __len__(self):
        """ Вернуть общее количество слов в словаре """
        return sum(len(i) for i in self.list_of_days)

    def __getitem__(self, item):
        """ Вернуть WordsPerDay элемент с датой, равной item
            Если день с датой пуст – вернуть пустой WordsPerDay-объект
            :param item: дата строкой, DATE-объектом или срезом
        """
        assert isinstance(item, DATE) or isinstance(item, str) or isinstance(item, slice), \
            f"Wrong item: '{item}', func – Vocabulary.__getitem__"

        # TODO: it could be some troubles after empty days removing
        if isinstance(item, slice):
            start = str_to_date(item.start) if item.start is not None else item.start
            stop = str_to_date(item.stop) if item.stop is not None else item.stop

            if start is not None and stop is not None and start > stop:
                raise ValueError(f"Start '{start}' cannot be more than stop '{stop}'")
            if start is not None and (start < self.begin() or start > self.end()):
                raise ValueError(f"Wrong start: '{start}'")
            if stop is not None and (stop > self.end() or stop < self.begin()):
                raise ValueError(f"Wrong stop: '{stop}'")

            begin = start if start is None else list(map(lambda x: x.get_date() == start, self.list_of_days)).index(True)
            end = stop if stop is None else list(map(lambda x: x.get_date() == stop, self.list_of_days)).index(True)
            # TODO: нет инита по списку
            # res = self.__class__()
            # res.list_of_days = self.list_of_days[begin:end]
            # return res

        item = str_to_date(item)

        assert item in self, \
            f"Wrong date: '{item}' func – Vocabulary.__getitem__"
        # дата в словаре, но её нет в списке дат – день пуст
        if item not in self.get_date_list():
            return WordsPerDay([], item)

        return list(filter(lambda x: item == x.get_date(), self.list_of_days))[0]

    def __call__(self, item, **kwargs):
        """ Вернуть все найденные слова строкой;
            Параметры поиска:
                1. by_def – искать в определениях; если в искомом элементе есть хоть
                    один русский символ – также искать в определениях;
        """
        assert (isinstance(item, str) or isinstance(item, Word)) and len(item) > 1, \
            f"Wrong word '{item}', func – Vocabulary.__call__"

        word = item.lower().strip() if isinstance(item, str) else item.word

        if 'by_def' in kwargs and kwargs['by_def']:
            word = f" – {word}"
        try:
            res = [f"{date}: {S_TAB.join(map(lambda x: up_word_in_def(str(x), word), words))}"
                   for date, words in self.search(word).items()]
            return '\n'.join(res)
        except Exception as trouble:
            return f"Word '{word}' not found, {trouble}"

    def __str__(self):
        """ Вернуть все дни и информацию о словаре """
        return f"{self.info()}\n{DIVIDER}\n" + \
               f"\n{DIVIDER}\n\n".join(map(str, self.list_of_days))

    def __bool__(self):
        """ Стандартный списочный """
        return bool(self.list_of_days)

    def __iter__(self):
        """ Стандартный списочный """
        return iter(self.list_of_days)

    def __hash__(self):
        return hash(sum(hash(i) for i in self.list_of_days))


def backup_everything(vocab, s_exams, c_exams):
    """ Backup: main vocabulary, self examples
        corpus example, repeat log
    """
    assert isinstance(vocab, Vocabulary) and vocab, \
        f"Wrong vocabulary: '{vocab}', func – backup_everything"
    assert isinstance(s_exams, SelfExamples) and s_exams, \
        f"Wrong self_examples: '{s_exams}', func – backup_everything"
    assert isinstance(c_exams, CorpusExamples) and c_exams, \
        f"Wrong corpus_examples: '{c_exams}', func – backup_everything"

    vocab.backup()
    print()
    s_exams.backup()
    print()
    c_exams.backup()
    print()
    backup_repeat_log()


# TODO: избегать флаги как параметры функции
# TODO: траблы с hint там, где есть собственные примеры
# TODO: сделать все функции выполняющими только одну поставленную задачу
# TODO: в поля классов добавить _ в начало имени
# TODO: проверка валидности введённого значения при restore
# TODO: ловля вообще всех ошибок в backup_setup
# TODO: функции отсортировать по модулям
# TODO: если после поимки исключения задача ф-ции не выполняется,
#  то печатать 'Terminating...'
