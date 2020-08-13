__all__ = (
    'create_pdf', 'visual_info', 'create_docx'
)

import os
from pathlib import Path
from typing import List, Dict

import comtypes.client
import docx
import xlsxwriter
from docx.shared import Pt

import src.main.common_funcs as comm_func
import src.main.constants as const

FONT_NAME = 'Avenir Next Cyr'
FONT_SIZE = 16


def create_docx(f_path: str or Path,
                content: List,
                header: str = 'General') -> None:
    """ Create docx file with strable content.
    If there's a file with tha same name, don't create the new one.

    Format:
    Enumerate and capitalize strings.

    :param f_path: str or Path, file name.
    :param content: list, strable objects to dump.
    :param header: str, doc header. 'General' by default.
    """
    if not all(hasattr(i, '__str__') for i in content):
        raise TypeError("List items must be strable")

    f_path = Path(f_path)

    if f_path.exists():
        raise FileExistsError(f"Doc '{f_path}' still exists")

    _docx = docx.Document()

    # font and its properties
    doc_style = _docx.styles['Normal']
    _font = doc_style.font
    _font.name = FONT_NAME
    _font.size = Pt(FONT_SIZE)

    _docx.add_heading(f"{header}", 0)
    for num, item in enumerate(content, 1):
        new_paragraph = _docx.add_paragraph()
        new_paragraph.style = doc_style

        new_paragraph.add_run(f"{num}. ").bold = True
        item = str(item)
        new_paragraph.add_run(f"{item[0].upper()}{item[1:]}")

    _docx.save(f_path)


def create_pdf(f_path: str or Path,
               content: List) -> None:
    """ Create a PDF file vie docx mediator.

    :param f_path: str or Path, path to file.
    :param content: list, strable objects.
    """
    f_path = Path(f_path)
    if f_path.exists():
        raise FileExistsError(f"PDF '{f_path}' still exists")

    # file mediator
    date = comm_func.today(const.DATEFORMAT)
    mediator_path = const.PDF_FOLDER / f"temp_{len(content)}_{date}.docx"
    try:
        create_docx(mediator_path, content)
    except FileExistsError:
        pass
    except Exception:
        raise

    full_mediator_path = Path.cwd() / mediator_path
    full_pdf_path = Path.cwd() / f_path

    word_client = comtypes.client.CreateObject('Word.Application')
    mediator = word_client.Documents.Open(full_mediator_path)
    mediator.SaveAs(full_pdf_path, FileFormat=17)

    mediator.Close()
    word_client.Quit()

    os.system(f'del "{full_mediator_path}"')


def visual_info(f_path: str,
                content: Dict,
                **kwargs) -> None:
    """ Create Excel file with graphic.
    By x – dict's keys, by y – dict's values.

    :keyword sheet_name: str, sheet name. 'Graphic' by default.
    :keyword chart_type: str, graphic type: line, area, bar, column etc.
     'line' by default.
    :keyword chart_title: str, graphic header. 'Title' by default.
    :keyword font_name: str, font name. 'Avenir Next Cyr' by default.
    :keyword x_axis_name: str, X axis name. 'Keys' by default.
    :keyword y_axis_name: – str, Y axis name. 'Values' by default.
    :keyword wb_title: – str, doc header. 'Title' by default.
    :keyword author: str, doc author. 'Author' by default.
    :keyword total: – добавлять ли в конце сумму всех значений; True by default.

    :param content: dict, pairs to write.
    :param f_path: str, file name.
    :return None.
    """
    if not isinstance(content, dict):
        raise TypeError(f"Content must be a dict, but {type(content)} given")
    if not all(isinstance(i, (int, float)) for i in content.values()):
        raise TypeError(f"Dict keys must be int or float")

    sheet_name = kwargs.pop('sheet_name', 'Graphic')
    chart_type = kwargs.pop('chart_type', 'line')
    chart_title = kwargs.pop('chart_title', 'Title')
    font = kwargs.pop('font_name', FONT_NAME)
    x_axis_name = kwargs.pop('x_axis_name', 'Keys')
    y_axis_name = kwargs.pop('y_axis_name', 'Values')
    wb_title = kwargs.pop('wb_title', 'Title')
    author = kwargs.pop('author', 'Author')

    f_path = Path(f_path)
    if f_path.exists():
        raise ValueError(f"File {f_path} still exists")

    wb = xlsxwriter.Workbook(f_path)
    wb.set_properties({
        'title': wb_title,
        'author': author,
        'comments': "Created with Python and XlsxWriter"
    })

    cell_format = wb.add_format({
        'size': 16,
        'font': font,
        'align': "vcenter"
    })

    work_sheet = wb.add_worksheet(sheet_name)
    work_sheet.set_column('A:A', 17)

    enum_cont = [
        (num, (key, val))
        for num, (key, val) in enumerate(content.items())
    ]
    for row, (key, val) in enum_cont:
        work_sheet.write(row, 0, str(key), cell_format)
        work_sheet.write(row, 1, val, cell_format)

    row += 1

    work_sheet.write(row, 0, "Total:", cell_format)
    work_sheet.write(row, 1, f"=SUM(B1:B{row})", cell_format)

    _chart = wb.add_chart({
        'type': chart_type
    })

    _chart.set_title({
        'name': chart_title,
        'name_font': {
            'name': font,
            'color': "black",
            'size': 16
        },
    })

    _chart.add_series({
        'values': f"={sheet_name}!B1:B{row}",
        'categories': f"={sheet_name}!A1:A{row}",
        'line': {
            'color': "orange"
        },
    })

    _chart.set_legend({
        'none': True
    })

    _chart.set_x_axis({
        'name': x_axis_name,
        'name_font': {
            'name': font,
            'italic': True,
            'size': 16
        },
        'num_font': {
                'name': font,
                'italic': True,
                'size': 14
            }
    })
    _chart.set_y_axis({
        'name': y_axis_name,
        'name_font': {
            'name': font,
            'italic': True,
            'size': 16
        },
        'num_font': {
            'name': font,
            'italic': True,
            'bold': True,
            'size': 14
         }
    })

    _chart.set_size({
        'width': 1280,
        'height': 520
    })

    work_sheet.insert_chart('C1', _chart)

    wb.close()